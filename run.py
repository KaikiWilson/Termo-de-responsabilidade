from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from getData import *
from privateData import *


document = Document()

document.add_picture('LogoBravaBrazil.png', width=Inches(2))
lastParagraph = document.paragraphs[-1] 
lastParagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_heading('DEPARTAMENTO DE INFORMÁTICA', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_heading('Termo de responsabilidade pela guarda e uso do equipamento de trabaho', level=3).alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph(f'Declaro para os devidos fins que no dia ')

p.add_run(f'{day} de {week} de {year} ').bold = True

p.add_run('recebi da empresa ').font.size = Pt(8.5)
p.add_run({empresa.nameEmpresa}).bold = True
p.add_run(f', CNPJ: {empresa.CNPJempresa}, o(s) equipamento(s) citado(s) neste termo, estando estes em perfeito estado de uso, e sendo destinado ao uso exclusivo no exercício das minhas funções com a empresa acima mencionada.').font.size = Pt(8.5)

content = document.add_paragraph()
contentRun = content.add_run('Comprometo-me a zelar por sua guarda, conservação e devolução futura. Nas condições de liberação, comprometo-me a substituí-lo em caso de perda e danos graves e/ou irreparáveis, que possam ocasionar a inutilização futura; Por idêntico ou similar (indicado pela BRAVA BRAZIL S/A).').font.size = Pt(8.5)

content2 = document.add_paragraph()
content2Run = content2.add_run('Enquanto em posse do equipamento, o solicitante fica ciente de que o equipamento estará sujeito a inspeções sem prévio aviso.').font.size = Pt(8.5)


p2 = document.add_paragraph()
run = p2.add_run('Identificação do solicitante:')
run.font.size = Pt(10)
p2.style = document.styles['Title']
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

listName = document.add_paragraph(f'Nome: ', style='List Bullet')
listNameRun = listName.add_run({requestingUser.name})
listNameRun.bold = True
listNameRun.font.size = Pt(9)

listCPF = document.add_paragraph(f'CPF: ', style='List Bullet')
listCPFRun = listCPF.add_run({requestingUser.cpf})
listCPFRun.bold = True
listCPFRun.font.size = Pt(9)


listEmail = document.add_paragraph(f'Email: ', style='List Bullet')
listEmailRun = listEmail.add_run(requestingUser.email)
listEmailRun.bold = True
listEmailRun.font.size = Pt(9)

listPhone = document.add_paragraph(f'Telefone: ', style='List Bullet')
listPhoneRun = listPhone.add_run(f'+55 (85) {requestingUser.phoneNumber[:5]}-{requestingUser.phoneNumber[-4:]}')
listPhoneRun.bold = True
listPhoneRun.font.size = Pt(9)


p3 = document.add_paragraph()
run = p3.add_run('Identificação do(s) equipamento(s):')
font = run.font
font.name = 'Calibri'
font.size = Pt(10)
p3.style = document.styles['Title']
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

listName = document.add_paragraph(f'', style='List Number')
listNameRun = listName.add_run('equipamento 1')
listNameRun.bold = True
listNameRun.font.size = Pt(9)

listCPF = document.add_paragraph(f'', style='List Number')
listCPFRun = listCPF.add_run('equipamento 2')
listCPFRun.bold = True
listCPFRun.font.size = Pt(9)
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()

document.add_paragraph('___________________________________' + '                                                          ' + '___________________________________')

document.add_paragraph()

footer = document.add_paragraph()
footerRun = footer.add_run(f'{empresa.nameEmpresa}\n{empresa.CNPJempresa}')
footerRun.font.size = Pt(6.5)
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.save('Termo de responsabilidade - setor de TI.docx')